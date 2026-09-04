import json
import re
from io import BytesIO

import streamlit as st
from google import genai
from google.genai import types
from pypdf import PdfReader
from docx import Document


APP_TITLE = "ResumeLens — ATS Resume Analyzer"
MODEL_NAME = "gemini-2.5-flash"


def extract_pdf_text(file_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages).strip()


def extract_docx_text(file_bytes: bytes) -> str:
    document = Document(BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs if p.text.strip()]

    # Also capture text from tables, which are common in resumes.
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts).strip()


def extract_resume_text(uploaded_file) -> str:
    data = uploaded_file.getvalue()
    suffix = uploaded_file.name.lower().rsplit(".", 1)[-1]

    if suffix == "pdf":
        return extract_pdf_text(data)
    if suffix == "docx":
        return extract_docx_text(data)
    raise ValueError("Unsupported file type. Please upload a PDF or DOCX resume.")


def clean_json_response(text: str) -> dict:
    """Handle normal JSON plus occasional Markdown code fences from the model."""
    text = text.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*```$", "", text)

    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("Gemini returned an invalid JSON response.")
    return json.loads(text[start:end + 1])


def analyze_resume(resume_text: str, target_role: str, api_key: str) -> dict:
    client = genai.Client(api_key=api_key)

    prompt = f"""
You are an expert ATS resume evaluator and career-document reviewer.

Analyze the resume below for the target role: "{target_role or "General professional role"}".

Important:
- This is an ATS-style estimate, not a guarantee of any employer's actual ATS result.
- Do not invent jobs, education, skills, achievements, metrics, dates, or certifications.
- Evaluate only what is present in the supplied resume.
- Give practical improvements that the applicant can actually make.
- Keep the advice professional and concise.
- Never judge the applicant based on age, gender, race, religion, disability, nationality, photo, or other protected/personal characteristics.

Return ONLY valid JSON matching this structure:
{{
  "ats_score": 0,
  "score_label": "Weak|Needs Work|Good|Strong|Excellent",
  "summary": "2-4 sentence overview",
  "category_scores": {{
    "keyword_alignment": 0,
    "format_structure": 0,
    "experience_impact": 0,
    "skills_relevance": 0,
    "clarity_readability": 0
  }},
  "strengths": ["...", "...", "..."],
  "improvements": [
    {{
      "priority": "High|Medium|Low",
      "area": "Keywords|Experience|Skills|Summary|Formatting|Projects|Education|Other",
      "issue": "...",
      "recommendation": "...",
      "example": "Optional example wording based only on facts already in the resume, or empty string"
    }}
  ],
  "missing_or_weak_sections": ["..."],
  "ats_keywords_to_consider": ["..."],
  "final_action_plan": ["...", "...", "..."]
}}

Scoring:
- ats_score must be an integer from 0 to 100.
- category_scores must each be integers from 0 to 100.
- Consider standard ATS concerns: readable structure, conventional headings, keyword alignment with the target role, measurable impact, skills coverage, consistency, and clarity.
- Do not penalize a resume simply for being visually plain.

Target role:
{target_role or "General professional role"}

Resume:
---BEGIN RESUME---
{resume_text[:50000]}
---END RESUME---
"""

    response = client.models.generate_content(
        model=MODEL_NAME,
        contents=prompt,
        config=types.GenerateContentConfig(
            temperature=0.2,
            response_mime_type="application/json",
        ),
    )

    result = clean_json_response(response.text)

    # Basic validation so malformed model output does not break the UI.
    score = int(result.get("ats_score", 0))
    result["ats_score"] = max(0, min(100, score))

    categories = result.get("category_scores", {})
    for key in (
        "keyword_alignment",
        "format_structure",
        "experience_impact",
        "skills_relevance",
        "clarity_readability",
    ):
        categories[key] = max(0, min(100, int(categories.get(key, 0))))
    result["category_scores"] = categories

    return result


def render_score(score: int) -> None:
    st.metric("Estimated ATS Score", f"{score}/100")
    st.progress(score / 100)

    if score >= 85:
        st.success("Excellent — the resume is well aligned and readable.")
    elif score >= 70:
        st.success("Strong — a few targeted improvements could make it better.")
    elif score >= 55:
        st.warning("Good foundation — several improvements are recommended.")
    elif score >= 40:
        st.warning("Needs work — ATS alignment and/or resume clarity can be improved.")
    else:
        st.error("Weak — substantial improvements are recommended.")


def main():
    st.set_page_config(
        page_title=APP_TITLE,
        page_icon="📄",
        layout="wide",
    )

    st.title("📄 ResumeLens")
    st.caption("Upload a resume to get an estimated ATS score and practical improvement suggestions.")

    with st.sidebar:
        st.header("Settings")
        api_key = st.text_input(
            "Gemini API key",
            type="password",
            help="For local use. On Streamlit Community Cloud, store this in Secrets instead of hard-coding it.",
        )

        if not api_key:
            try:
                api_key = st.secrets["GEMINI_API_KEY"]
            except Exception:
                api_key = ""

        target_role = st.text_input(
            "Target job title",
            placeholder="e.g. Python Developer",
        )

        st.info(
            "Tip: A target role makes keyword-alignment scoring more useful. "
            "The app does not guarantee a real employer ATS score."
        )

    uploaded_file = st.file_uploader(
        "Upload your resume",
        type=["pdf", "docx"],
        accept_multiple_files=False,
    )

    if not uploaded_file:
        st.markdown(
            """
            ### What you'll get
            - **Estimated ATS score** out of 100
            - Category-by-category scoring
            - Resume strengths
            - High/medium/low priority improvements
            - ATS keywords to consider
            - A simple final action plan
            """
        )
        return

    if uploaded_file.size > 10 * 1024 * 1024:
        st.error("Please upload a resume smaller than 10 MB.")
        return

    try:
        resume_text = extract_resume_text(uploaded_file)
    except Exception as exc:
        st.error(f"Could not read this resume: {exc}")
        return

    if len(resume_text.strip()) < 100:
        st.error(
            "Very little text could be extracted. If this is a scanned/image-only PDF, "
            "please use a text-based PDF or DOCX version."
        )
        return

    with st.expander("Preview extracted resume text"):
        st.text(resume_text[:12000])

    if st.button("🔎 Analyze Resume", type="primary", use_container_width=True):
        if not api_key:
            st.error(
                "Add your Gemini API key in the sidebar, or configure GEMINI_API_KEY "
                "in Streamlit Secrets."
            )
            return

        with st.spinner("Analyzing your resume with Gemini Flash..."):
            try:
                result = analyze_resume(resume_text, target_role, api_key)
            except Exception as exc:
                st.error(
                    "Analysis failed. Check your Gemini API key, model/API availability, "
                    "and the Streamlit logs for details."
                )
                st.exception(exc)
                return

        st.divider()
        col1, col2 = st.columns([1, 2])

        with col1:
            render_score(result["ats_score"])

        with col2:
            st.subheader("Summary")
            st.write(result.get("summary", ""))

        st.subheader("Category scores")
        scores = result["category_scores"]
        cols = st.columns(5)
        labels = [
            ("Keyword alignment", scores["keyword_alignment"]),
            ("Format & structure", scores["format_structure"]),
            ("Experience impact", scores["experience_impact"]),
            ("Skills relevance", scores["skills_relevance"]),
            ("Clarity", scores["clarity_readability"]),
        ]
        for col, (label, value) in zip(cols, labels):
            col.metric(label, f"{value}/100")

        left, right = st.columns(2)

        with left:
            st.subheader("✅ Strengths")
            for item in result.get("strengths", []):
                st.markdown(f"- {item}")

            st.subheader("📌 Missing or weak sections")
            missing = result.get("missing_or_weak_sections", [])
            if missing:
                for item in missing:
                    st.markdown(f"- {item}")
            else:
                st.write("No major missing sections were identified.")

        with right:
            st.subheader("🔑 ATS keywords to consider")
            keywords = result.get("ats_keywords_to_consider", [])
            if keywords:
                st.write(", ".join(keywords))
            else:
                st.write("No additional keywords were suggested.")

        st.subheader("🛠️ Improvement plan")
        improvements = result.get("improvements", [])
        for index, item in enumerate(improvements, start=1):
            priority = item.get("priority", "Medium")
            area = item.get("area", "Other")
            with st.expander(f"{index}. [{priority}] {area}: {item.get('issue', '')}"):
                st.write(item.get("recommendation", ""))
                example = item.get("example", "")
                if example:
                    st.markdown("**Possible wording:**")
                    st.info(example)

        st.subheader("🚀 Final action plan")
        for item in result.get("final_action_plan", []):
            st.markdown(f"- {item}")

        st.caption(
            "Privacy note: this app sends the extracted resume text to Gemini for analysis. "
            "Do not upload documents containing information you do not want sent to the configured API."
        )


if __name__ == "__main__":
    main()
